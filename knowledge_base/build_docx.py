"""Convert the markdown knowledge base into real Word (.docx) files.

The MSAT Agent's grounding documents are managed as Word files; this script is
the one-time (re-runnable) generator that turns each ``NN_*.md`` into a matching
``NN_*.docx`` with proper Word styles (Heading 1/2/3, tables, bullet lists) so
the reference viewer can render a genuine Word-style preview.

Run:  python knowledge_base/build_docx.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

KB = Path(__file__).parent
INLINE = re.compile(r"[`*]")          # strip markdown bold / inline-code markers


def _clean(s: str) -> str:
    return INLINE.sub("", s).strip()


def _add_meta(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(_clean(text))
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x6B, 0x6B, 0x6B)


def _split_row(line: str) -> list[str]:
    cells = line.strip().strip("|").split("|")
    return [_clean(c) for c in cells]


def _is_sep_row(line: str) -> bool:
    return bool(re.match(r"^\s*\|?[\s:|-]+\|?\s*$", line)) and "-" in line


def md_to_docx(md_path: Path) -> Path:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # table: a run of pipe rows
        if stripped.startswith("|") and i + 1 < len(lines) and _is_sep_row(lines[i + 1]):
            header = _split_row(stripped)
            rows: list[list[str]] = []
            j = i + 2
            while j < len(lines) and lines[j].strip().startswith("|"):
                rows.append(_split_row(lines[j]))
                j += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            for c, txt in enumerate(header):
                table.rows[0].cells[c].text = txt
            for r in rows:
                cells = table.add_row().cells
                for c in range(len(header)):
                    cells[c].text = r[c] if c < len(r) else ""
            i = j
            continue

        # headings
        m = re.match(r"^(#{1,3})\s+(.*)", stripped)
        if m:
            level = len(m.group(1))
            doc.add_heading(_clean(m.group(2)), level=level)
            i += 1
            continue

        # bullet list
        if stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(_clean(stripped[2:]), style="List Bullet")
            i += 1
            continue

        # metadata line (Document ID / Revision …)
        if stripped.startswith("**Document ID") or "Revision:" in stripped:
            _add_meta(doc, stripped)
            i += 1
            continue

        doc.add_paragraph(_clean(stripped))
        i += 1

    out = md_path.with_suffix(".docx")
    doc.save(out)
    return out


def main() -> None:
    made = []
    for md in sorted(KB.glob("*.md")):
        if md.name.lower() == "readme.md":
            continue
        made.append(md_to_docx(md).name)
    print(f"generated {len(made)} docx files:")
    for n in made:
        print("  -", n)


if __name__ == "__main__":
    main()
